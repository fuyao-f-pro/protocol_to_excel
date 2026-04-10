"""
parsers/word_parser.py — Word (.docx) 解析
"""
import os
from typing import List, Dict, Optional
import pandas as pd
from docx import Document as DocxDoc
from models import ParsedDocument
class WordParser:
    def __init__(self, filepath: str):
        self.filepath = filepath
        self.filename = os.path.basename(filepath)
    def parse(self) -> ParsedDocument:
        doc = ParsedDocument(filename=self.filename)
        docx = DocxDoc(self.filepath)
        # 段落 → 纯文本
        text_parts = [p.text.strip() for p in docx.paragraphs if p.text.strip()]
        doc.raw_text = "\n".join(text_parts)
        # 表格
        tables = []
        for idx, table in enumerate(docx.tables):
            df = self._table_to_df(table)
            if df is not None:
                title = self._find_title(docx, table)
                tables.append({"index": idx, "title": title, "data": df})
        doc.raw_tables = tables
        return doc
    @staticmethod
    def _table_to_df(table) -> Optional[pd.DataFrame]:
        rows = []
        for row in table.rows:
            cells = [c.text.replace("\n", " ").strip() for c in row.cells]
            rows.append(cells)
        if len(rows) < 2:
            return None
        # 去除合并单元格产生的连续重复
        cleaned = []
        for row in rows:
            new_row = []
            prev = object()
            for c in row:
                new_row.append("" if c == prev and c != "" else c)
                prev = c
            cleaned.append(new_row)
        max_cols = max(len(r) for r in cleaned)
        cleaned = [r + [""] * (max_cols - len(r)) for r in cleaned]
        header = cleaned[0]
        data = cleaned[1:]
        try:
            df = pd.DataFrame(data, columns=header).drop_duplicates()
            return df if len(df) > 0 else None
        except Exception:
            return pd.DataFrame(data)
    @staticmethod
    def _find_title(docx, target_table) -> str:
        """找到表格前最近的段落文本作为标题"""
        last_text = ""
        for elem in docx.element.body:
            tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag
            if tag == "p":
                # 是段落
                for para in docx.paragraphs:
                    if para._element is elem:
                        t = para.text.strip()
                        if t:
                            last_text = t
                        break
            elif tag == "tbl":
                # 是表格
                if elem is target_table._element:
                    return last_text
        return last_text
